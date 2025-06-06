'''
Author:Omar Usmani (Omar.Usmani@TNO.nl).
This module is a cookbook with useful auxiliary functions.
They are listed in the following documentation:
https://tno.github.io/ETS_CookBook/
'''

import collections
import datetime
import functools
import math
import os
import sqlite3
import time
import tomllib
import typing as ty
import zipfile

import box
import dash
import docx
import docx.document
import geopandas as gpd
import matplotlib
import matplotlib.axes
import matplotlib.colors
import matplotlib.figure
import matplotlib.projections
import matplotlib.pyplot as plt
import numpy as np
import openpyxl
import openpyxl.worksheet
import openpyxl.worksheet.cell_range
import openpyxl.worksheet.table
import pandas as pd
import plotly
import plotly.graph_objects as go
import requests
import xarray as xr


def check_if_folder_exists(folder_to_check: str) -> None:
    '''
    Checks if a folder exists. If it does not, it creates it.
    This way, users can set up a new (sub-)folder in the configuration file
    without having to ensure that it already exits or create it.
    '''

    # We check if the output folder exists.
    if not os.path.exists(folder_to_check):
        # If it doesn't, we create it
        os.makedirs(folder_to_check)


def parameters_from_TOML(parameters_file_name: str) -> box.Box:
    '''
    Reads a TOML parameters file name and returns a parameters Box.
    '''

    with open(parameters_file_name, mode='rb') as parameters_file:
        parameters: box.Box = box.Box(tomllib.load(parameters_file))

    return parameters


def reference_scale(
    number_list: list[float], digit_shift: int = 0
) -> list[float]:
    '''
    This function takes a list of numbers and returns a scale
    (lower and upper boundary) they are in.
    The digit shift parameter tells us on which digit we need to
    focus. The default is 0, so the upper boundary of 53.57 will be 60
    by default, but 54 if the digit shift is 1 (thus focussing on the 3 part).
    This can for example be useful to determine the plotting area of a dataset
    (x-axis boundaries).
    '''

    number_list_boundaries: list[float] = [
        min(number_list),
        max(number_list),
    ]
    boundary_powers_of_ten: list[int] = [
        (
            int(math.log10(abs(number))) - digit_shift
            # The first term gives us the power of ten of the highest digit,
            # which we shift with the digit_shift parameter
            if number != 0
            else 0
        )
        for number in number_list_boundaries
    ]

    dividers: list[float] = [
        math.pow(10, power) for power in boundary_powers_of_ten
    ]

    lower_scale: float = (
        math.floor((number_list_boundaries[0]) / dividers[0]) * dividers[0]
    )

    upper_scale: float = (
        math.ceil((number_list_boundaries[1]) / dividers[1]) * dividers[1]
    )

    return [lower_scale, upper_scale]


def dataframe_from_Excel_table_name(
    table_name: str, Excel_file: str, load_data_only: bool = True
) -> pd.DataFrame:
    '''
    This function looks up a given table name in an Excel file
    and returns a DataFrame containing the values of that table.
    Note that if the name does not exist (or is spelled wrongly (it's
    case-sensitive), the function will crash).
    The optional load_data_only parameter puts values in the table if set to
    True (its default value. A False value loads formulas)
    '''

    source_workbook: openpyxl.Workbook = openpyxl.load_workbook(
        Excel_file, data_only=load_data_only
    )

    # We need to look up the worksheet
    # table_worksheet = ''
    for worksheet in source_workbook.worksheets:
        if table_name in worksheet.tables:
            table_worksheet = worksheet

    table: openpyxl.worksheet.table.Table = table_worksheet.tables[table_name]
    table_range: tuple = table_worksheet[table.ref]

    table_entries: list[list] = [
        [cell_entry.value for cell_entry in row_entry]
        for row_entry in table_range
    ]

    table_headers: list[str] = table_entries[0]
    table_values: list[ty.Any] = table_entries[1:]

    # We need to go thhrough a dictionary, as passing the values directly
    # leads to duplicates with no index, for some reason
    table_dictionary: dict[str, ty.Any] = {}
    for header_index, header in enumerate(table_headers):
        values_for_header: list[ty.Any] = [
            table_row[header_index] for table_row in table_values
        ]
        table_dictionary[header] = values_for_header

    table_dataframe: pd.DataFrame = pd.DataFrame.from_dict(table_dictionary)

    return table_dataframe


def dataframe_to_Excel(
    dataframe_to_append: pd.DataFrame, Excel_workbook: str, my_sheet: str
) -> None:
    '''
    This function takes a DataFrame and puts it into a new sheet in
    an Excel workbook. If the sheet already exists, it will replace it.
    If the Excel workbook does not exist, the function creates it.
    '''

    if not os.path.isfile(Excel_workbook):
        new_workbook: openpyxl.Workbook = openpyxl.Workbook()
        new_workbook.save(Excel_workbook)

    with pd.ExcelWriter(
        Excel_workbook, mode='a', engine='openpyxl', if_sheet_exists='replace'
    ) as writer:
        dataframe_to_append.to_excel(writer, sheet_name=my_sheet)


def get_extra_colors(color_definitions: box.Box) -> pd.DataFrame:
    '''
    This function gets user-defined extra colors from a Box..
    This Box contains the names of the colors, and their RGB values
    (from 0 to 255). The function returns a DataFrame with
    color names as index, and their RGB codes (between 0 and 1)
    as values.
    '''

    extra_colors: pd.DataFrame = pd.DataFrame(columns=['R', 'G', 'B'])
    for color_definition in color_definitions:
        extra_colors.loc[color_definition] = color_definitions[
            color_definition
        ]
    extra_colors = extra_colors / 255

    return extra_colors


def get_rgb_from_name(
    color_name: str, color_definitions: box.Box
) -> list[float]:
    '''
    This function takes a color name and returns its RGB values (0 to 1).
    If the color name is in the extra colors, then, we use
    the values given.
    If it is a matplotlib color, then we use the matplotlib function.
    '''
    extra_colors: pd.DataFrame = get_extra_colors(color_definitions)

    if color_name in extra_colors.index.values:
        return list(extra_colors.loc[color_name].values)
    else:
        return list(matplotlib.colors.to_rgb(color_name))


def rgb_color_list(
    color_names: list[str], color_definitions: box.Box
) -> list[list[float]]:
    '''
    Gets a list of RGB codes for a list of color names.
    '''
    rgb_codes: list[list[float]] = [
        get_rgb_from_name(color_name, color_definitions)
        for color_name in color_names
    ]

    return rgb_codes


def register_color_bars(
    color_bar_definitions: box.Box, color_definitions: box.Box
) -> None:
    '''
    This function reads the user-defined color bars in Box
    (names for the bars, and a list of the colors they contain, with the
    first being at the bottom of the bar, and the last at the top, and the
    others in between). It then creates the color bars and stores them
    in the list of available color maps.
    '''

    # This dictionary stores the dictionaries for each color bar
    color_bar_dictionary: dict = {}
    # These colors are the three base keys of each bar color dictionary
    # Each dictionary contains a tuple of tuples for each base colors
    # Each of these sub-tuples cotains a step (between 0 and 1),
    # and a tone of the basic color in question (red, green, or blue)
    # This is repeated. The second value can be different
    # if cretaing discontinuities.
    # See https://matplotlib.org/stable/gallery/color/custom_cmap.html
    # for details
    base_colors_for_color_bar: list[str] = ['red', 'green', 'blue']

    # We fill the color bar dictionary
    for color_bar in color_bar_definitions:
        # We read the color list
        color_bar_colors: list[str] = color_bar_definitions[color_bar]
        # We set the color steps, based on the color list
        color_steps: np.ndarray = np.linspace(0, 1, len(color_bar_colors))

        color_bar_dictionary[color_bar] = {}
        for base_color_index, base_color in enumerate(
            base_colors_for_color_bar
        ):
            # We create a list of entries for that base color
            # It is a list so that we can append,
            # but we will need to convert it to a tuple
            base_color_entries = []
            for color_step, color_bar_color in zip(
                color_steps, color_bar_colors
            ):
                # We get the tone by getting the RGB values of the
                # color bar color and taking the corresponding base index
                color_bar_color_tone = get_rgb_from_name(
                    color_bar_color, color_definitions
                )[base_color_index]

                base_color_entries.append(
                    # The subtuples consist of the color step
                    (
                        color_step,
                        # And the tone of the base color
                        # for the color corresponding
                        # to the step
                        color_bar_color_tone,
                        # This is repeated for continuous schemes
                        # See
                        # https://matplotlib.org/stable/gallery/color/custom_cmap.html
                        # for details
                        color_bar_color_tone,
                    )
                )

            # We now convert the list to a tuple and put it into
            # the dictionary
            color_bar_dictionary[color_bar][base_color] = tuple(
                base_color_entries
            )

    # We now add the color bars to the color maps

    for color_bar in color_bar_definitions:
        color_bar_to_register: matplotlib.colors.LinearSegmentedColormap = (
            matplotlib.colors.LinearSegmentedColormap(
                color_bar, color_bar_dictionary[color_bar]
            )
        )

        if color_bar_to_register.name not in matplotlib.pyplot.colormaps():
            matplotlib.colormaps.register(color_bar_to_register)


def get_season(time_stamp: datetime.datetime) -> str:
    '''
    This function takes a datetime timestamp and tells us in which season
    it is.
    '''
    # We take the date of the timestamp to avoid issues
    # during the transitions (where the timestamp would be
    # larger than the previous season's end, which is at midnight)
    date: datetime.datetime = datetime.datetime(
        time_stamp.year, time_stamp.month, time_stamp.day, 0, 0
    )
    seasons: list[tuple[str, tuple[datetime.datetime, datetime.datetime]]] = [
        (
            'winter',
            (
                datetime.datetime(date.year, 1, 1),
                datetime.datetime(date.year, 3, 20),
            ),
        ),
        (
            'spring',
            (
                datetime.datetime(date.year, 3, 21),
                datetime.datetime(date.year, 6, 20),
            ),
        ),
        (
            'summer',
            (
                datetime.datetime(date.year, 6, 21),
                datetime.datetime(date.year, 9, 22),
            ),
        ),
        (
            'fall',
            (
                datetime.datetime(date.year, 9, 23),
                datetime.datetime(date.year, 12, 20),
            ),
        ),
        (
            'winter',
            (
                datetime.datetime(date.year, 12, 21),
                datetime.datetime(date.year, 12, 31),
            ),
        ),
    ]
    for season, (start, end) in seasons:
        if start <= date <= end:
            result_season: str = season
    return result_season


def save_figure(
    figure: matplotlib.figure.Figure,
    figure_name: str,
    output_folder: str,
    dpi_to_use: int,
    file_formats: box.Box,
) -> None:
    '''
    This function saves a Matplolib figure to a number of
    file formats and an output folder.
    '''

    check_if_folder_exists(output_folder)

    for file_format in file_formats:
        if file_formats[file_format]:
            figure.savefig(
                f'{output_folder}/{figure_name}.{file_format}', dpi=dpi_to_use
            )


def save_dataframe(
    dataframe: pd.DataFrame,
    dataframe_name: str,
    groupfile_name: str,
    output_folder: str,
    dataframe_formats: box.Box,
) -> None:
    '''
    This function saves a pandas dataframe to a number of
    file formats and an output folder.

    Note that for some file types, you might need to install additional
    libraries.

    Also note that some formats will be saved into a group file that
    can contain several other dataframes (for example sheets into an Excel
    workbook or tables into an SQL database). dataframe_name will be the
    file name if the file format does not use group files. If the format does
    use a group fiule, then dataframe_name will be used for the sub-elements
    (sheets, tables, for example), and groupfile_name will be used for
    the file name (you can of course use the same value for both), and will be
    unused if the file format does not use group files.


    Bug to fix: XML does not accet a number to start names (or
    various case variations of xml), which must currently
    be handled by the user (who must avoid these).
    Other bug: Removing non-alphanumeric characters in the index does not seem
    to work (for xml and stata)

    gbq and orc outputs are not currently supported, as gbq is not
    a local file format, but a cloud-based one and orc does not seem to work
    with pyarrow (at least in Windows).
    '''

    check_if_folder_exists(output_folder)

    file_types: list[str] = [
        'csv',
        'json',
        'html',
        'latex',
        'xml',
        'clipboard',
        'excel',
        'hdf',
        'feather',
        'parquet',
        'stata',
        'pickle',
        'sql',
    ]
    # Note that pandas has a few more export formats that we skipped
    # orc is not supported in arrows (ar least on Windows)
    # https://stackoverflow.com/questions/58822095/no-module-named-pyarrow-orc
    # gbq is about Google cloud storage, not about local files
    # https://cloud.google.com/bigquery/docs/introduction
    # Note that clipboard does not produce a file,
    # but can still be used locally, so the function supports it.

    file_extensions: list[str] = [
        'csv',
        'json',
        'html',
        'tex',
        'xml',
        '',
        'xlsx',
        'h5',
        'feather',
        'parquet',
        'dta',
        'pkl',
        'sqlite3',
    ]

    # This determines if the dataframe is saved into its own file
    # or into a group file (such as a database or an Excel Workbook)
    is_groupfile_per_type: list[bool] = [
        False,
        False,
        False,
        False,
        False,
        False,
        True,
        True,
        False,
        False,
        False,
        False,
        True,
    ]

    file_functions: list[ty.Callable] = []
    for file_type, is_groupfile in zip(file_types, is_groupfile_per_type):

        # Some file formats require some extra processing, so
        # we need to forgo list compreshension and use
        # copies of the dataframe that we will modify

        dataframe_to_use: pd.DataFrame = dataframe.copy()

        if file_type == 'feather':
            # feather does not support serializing
            #  <class 'pandas.core.indexes.base.Index'> for the index;
            #   you can .reset_index() to make the index into column(s)
            dataframe_to_use = dataframe_to_use.reset_index()

        if (file_type in ['xml', 'stata']) and is_groupfile:
            # These file formats have issues with some characters
            # in column and index names
            # Note that for xml, the names cannot start with the letters
            # xaml (with all case variations) and must start with a letter
            # or underscore (replacement of such issues is not implemented
            # at the moment, so the function will fail unless you correct
            # that in your data).
            # Note that stata also has issues with too lonmg names
            #  (>32 characters), but the to_stata function manages this on
            # its own (by cutting any excess characters). As such, this
            # does not need to be corrected here
            # We only do this if the user wants to use these file types,
            # as it needs column headers to be strings to work,
            # but other file types don't need all this and thus
            # can still use column headers that aren't strings
            code_for_invalid_characters: ty.Literal['[^0-9a-zA-Z_.]'] = (
                '[^0-9a-zA-Z_.]'
            )

            dataframe_to_use.columns = dataframe_to_use.columns.str.replace(
                code_for_invalid_characters, '_', regex=True
            )

            if dataframe_to_use.index.name:
                # We only need to do the replacements if the index
                # has a name. This causes issues if we try to replace
                # things if the index does not have a name
                dataframe_to_use.index.name = (
                    dataframe_to_use.index.name.replace(' ', '_')
                )
            elif dataframe_to_use.index.names:
                # MultiIndex has to be treated sepaartely
                if dataframe_to_use.index.names[0] is not None:
                    dataframe_to_use.index.names = [
                        old_name.replace(' ', '_')
                        for old_name in dataframe_to_use.index.names
                    ]

        function_name: str = f'to_{file_type}'

        if file_type == 'latex':
            # In future versions `DataFrame.to_latex` is expected to
            # utilisethe base implementation of `Styler.to_latex` for
            # formatting and rendering. The arguments signature may
            # therefore change.
            # It is recommended instead to use `DataFrame.style.to_latex`
            # which also contains additional functionality.
            # function_name = f'Styler.to_{file_type}'
            file_functions.append(dataframe_to_use.style.to_latex)
        else:
            file_functions.append(getattr(dataframe_to_use, function_name))

    using_file_types: list[bool] = [
        dataframe_formats[file_type] for file_type in file_types
    ]

    for (
        file_type,
        file_extension,
        file_function,
        using_file_type,
        is_groupfile,
    ) in zip(
        file_types,
        file_extensions,
        file_functions,
        using_file_types,
        is_groupfile_per_type,
    ):
        if using_file_type:
            if is_groupfile:
                file_to_use: str = (
                    f'{output_folder}/{groupfile_name}.{file_extension}'
                )

                if file_type == 'hdf':
                    file_function(file_to_use, key=dataframe_name)
                elif file_type == 'excel':
                    # If we want to append a sheet to an Excel file
                    # instead of replacing the existing file, we need
                    # to use Excelwriter, but that gives an error if the
                    # file does not exist, so we need to check if the file
                    # exists
                    if os.path.exists(file_to_use):
                        writer_to_use: pd.ExcelWriter = pd.ExcelWriter(
                            file_to_use,
                            engine='openpyxl',
                            mode='a',
                            if_sheet_exists='replace',
                        )
                        with writer_to_use:
                            file_function(
                                writer_to_use, sheet_name=dataframe_name
                            )
                    else:
                        # If the file does not exist, we need to use the
                        # function, which is to_excel() with the file name,
                        # not with a writer
                        file_function(file_to_use, sheet_name=dataframe_name)

                elif file_type == 'sql':
                    with sqlite3.connect(file_to_use) as sql_connection:
                        file_function(
                            dataframe_name,
                            con=sql_connection,
                            if_exists='replace',
                        )

            else:
                file_to_use = (
                    f'{output_folder}/{dataframe_name}.{file_extension}'
                )
                file_function(file_to_use)


def put_dataframe_in_sql_in_chunks(
    source_dataframe: pd.DataFrame,
    sql_file: str,
    table_name: str,
    chunk_size: int,
    drop_existing_table: bool = True,
) -> None:
    '''
    This function takes a Dataframe and writes it into the table
    of an SQL database. It does so in chunks to avoid memory issues.
    The parameter drop_existing table tells us if we want to
    drop/overwrite the table if it exists (it is True by default).
    If set to False, the data will be appended (if the table exists).
    '''

    # We first need the total data/index length of the Dataframe
    data_length: int = len(source_dataframe.index)

    # We initialise the chunk boundaries and the sql connection
    chunk_start: int = 0
    chunk_end: int = 0

    with sqlite3.connect(sql_file) as sql_connection:
        if drop_existing_table:
            table_action: ty.Literal['fail', 'replace', 'append'] = 'replace'
        else:
            table_action = 'append'

        while chunk_end < data_length:
            chunk_end = min(chunk_start + chunk_size, data_length)
            # We select the corresponding chunk in the dataframe and
            # write it to the SQL database
            dataframe_chunk: pd.DataFrame = source_dataframe.iloc[
                chunk_start:chunk_end
            ]

            dataframe_chunk.to_sql(
                table_name, con=sql_connection, if_exists=table_action
            )

            chunk_start = chunk_end + 1
            # Subsequent additions append in all cases
            table_action = 'append'


def query_list_from_file(sql_file: str) -> list[str]:
    '''
    This returns a list of queries from an SQL file
    '''

    with open(sql_file) as script_file:
        sql_queries: list[str] = script_file.read().split(';')

    sql_queries.remove('')
    return sql_queries


def dataframes_from_query_list(
    query_list: list[str], sql_connection: sqlite3.Connection
) -> list[pd.DataFrame]:
    '''
    This returns a list of dataframes, each obtained from a query in the list
    '''
    dataframe_list: list[pd.DataFrame] = [
        pd.read_sql(sql_query, sql_connection) for sql_query in query_list
    ]

    return dataframe_list


def from_grib_to_dataframe(grib_file: str) -> pd.DataFrame:
    '''
    This function takes a grib file and converts it to a DataFrame.
    **Important note:**
    You need to have ecmwflibs installed for the grib converter to work.
    Installing xarray (and cfrgrib to have the right engine) is not enough!
    See:
    https://github.com/ecmwf/eccodes-python/issues/54#issuecomment-925036724
    '''
    grib_engine: str = 'cfgrib'

    source_data: xr.Dataset = xr.load_dataset(grib_file, engine=grib_engine)
    source_dataframe: pd.DataFrame = source_data.to_dataframe()

    return source_dataframe


def read_query_generator(
    quantities_to_display: str,
    source_table: str,
    query_filter_quantities: list[str] = [''],
    query_filter_types: list[str] = [''],
    query_filter_values: list = [''],
    # This can be a List of strings, or a nested list (see explanations)
) -> str:
    '''
    This function returns an sql read/select query string that can be used
    (for example) in Panda's read_sql.
    The input parameters are:
    - quantities_to_display: A string list of table column names
    (as strings, in
    single quotes), separated by commas. If the user
    wants all columns displayed, then they should use a '*'. If one (or more)
    of the column names have spaces, then the user needs to use f strings and
    double quotes, as in the following example:
    quantity_1 = 'Time'
    quantity_2 =  'Surveyed Area'
    quantity_2_with_quotes = f'"quantity_2"'
    quantities_to_display = f'{quantity_1}, {quantity_2_with_quotes}'
    This latter variable is the input for the function
    - source table is the name of the source table. Note that it has a similar
    need if the name has spaces, so use:
    source_table = f'"My Table"'
    as an input
    - query_filter_quantities: A list of strings each representing a column
    name the user wants to filter. Again, names with spaces require
    f strings and double quotes, so add:
    f'"Surveyed Area"' to your list of filter names
    - query_filter_types: This list (that has to be the same length as the
    above list of quantities)
    says which filter to use. Currently supported options are:
        - '='       (equal to)
        - '<'       (smaller than)
        - '>'       (larger than)
        - '!='      (not equal)
        - '<>'      (not equal)
        - '<='      (smalller or equal)
        - '>='      (larger or equal)
        - 'like'      (matches/ searches for a pattern)
        - 'between'   (between two  values)
        - 'in'        (to select multiple values for one or several columns)
    - query_filter_values: The comparison values used for the filter.
    The three special cases are:
        1) Like: This needs to be a double quote string (since it will be
        nested into a single-quote string) with percentage signs,
        such as '"%2020-05-08%"' for timestamps for May 8th, 2020
        2) Between: Provide the two  values  into a
        list. If the values arte strings that contain spaces,
        you need nested quotes, such as:
        ['"2020-05-08 00:00:00"','"2020-06-26 16:00:00"']
        3) In provide the two tuple values into a list.,
        e.g: [(52.1,4.9),(52.0,5.1)]

    '''

    query_filter: str = make_query_filter(
        query_filter_quantities, query_filter_types, query_filter_values
    )

    output_query: str = (
        f'select {quantities_to_display} from {source_table} '
        f'{query_filter};'
    )

    return output_query


def database_tables_columns(database: str) -> dict:
    '''
    Returns a dictionary with the tables of a database as keys and their
    columns as values.
    '''
    database_connection: sqlite3.Connection = sqlite3.connect(database)
    tables_query: str = 'select name from sqlite_master where type="table";'
    database_cursor: sqlite3.Cursor = database_connection.cursor()
    database_cursor.execute(tables_query)
    database_tables: list[str] = database_cursor.fetchall()
    tables_columns: dict[str, ty.Any] = {}
    for table in database_tables:
        table_name: str = table[0]
        # table_name = f'"{table[0]}"'
        table_cursor: sqlite3.Cursor = database_connection.execute(
            f'select * from "{table_name}"'
        )
        tables_columns[table_name] = [
            description[0] for description in table_cursor.description
        ]

    return tables_columns


def download_and_save_file(download_url: str, output_folder: str) -> None:
    '''
    Downloads a file from an URL and saves it. If the file is a zip file,
    the function extracts its contents.
    '''
    request_data: requests.Response = requests.get(download_url)
    file_name: str = download_url.split('/')[-1]

    with open(f'{output_folder}/{file_name}', 'wb') as output_file:
        output_file.write(request_data.content)

    if file_name.split('.')[-1] == 'zip':
        zip_data: zipfile.ZipFile = zipfile.ZipFile(
            f'{output_folder}/{file_name}'
        )
        for zip_info in zip_data.infolist():
            zip_data.extract(zip_info, path=output_folder)


def string_to_float(my_string: str) -> float:
    '''
    Converts strings to floats, and to zero if the string is not a float.
    '''
    try:
        my_output: float = float(my_string)
    except ValueError:
        my_output = 0.0

    return my_output


def get_map_area_data(map_parameters: box.Box) -> gpd.GeoDataFrame:
    '''
    This function gets and processes the area data and sets it
    into a DataFrame. It contains polygons/multipolygons
    (they are at a given granularity level, but also have references to higher
    levels).
    '''
    map_data_folder: str = map_parameters.map_data_folder
    # This file contains data at NUTS level 3. The reason for this is so that
    # we can remove the outer regions (such as Svalbard or French overseas
    # territories).
    area_data_file_name: str = map_parameters.area_data_file_name

    area_data: gpd.GeoDataFrame = gpd.read_file(
        f'{map_data_folder}/{area_data_file_name}'
    )

    # This is the list of regions to remove from the map.
    # These are the outer regions (such as Svalbard or French overseas
    # territories).
    general_exclusion_codes: list[str] = map_parameters.general_exclusion_codes

    # This removes the excluded areas. The ~ flips the boolean values
    # so that we keep the areas that are not in the exclusion list.
    area_data = area_data[~area_data['NUTS_ID'].isin(general_exclusion_codes)]

    return area_data


def get_map_borders(
    NUTS_level: int, map_parameters: box.Box
) -> gpd.GeoDataFrame:
    '''
    This function gets the borders/contours of regions at a specified NUTS
    level.
    '''
    map_data_folder: str = map_parameters.map_data_folder
    border_data_file_prefix: str = map_parameters.border_data_file_prefix
    border_data_file_suffix: str = map_parameters.border_data_file_suffix

    border_data_file: str = (
        f'{border_data_file_prefix}{NUTS_level}{border_data_file_suffix}'
    )

    border_data: gpd.GeoDataFrame = gpd.read_file(
        f'{map_data_folder}/{border_data_file}'
    )

    return border_data


def get_map_points(
    NUTS_level: int, map_parameters: box.Box
) -> gpd.GeoDataFrame:
    '''
    This function gets the points/labels of regions at a specified NUTS
    level.
    '''
    map_data_folder: str = map_parameters.map_data_folder
    points_data_file_prefix: str = map_parameters.points_data_file_prefix
    points_data_file_suffix: str = map_parameters.points_data_file_suffix

    points_data_file: str = (
        f'{points_data_file_prefix}{NUTS_level}{points_data_file_suffix}'
    )

    points_data: gpd.GeoDataFrame = gpd.read_file(
        f'{map_data_folder}/{points_data_file}'
    )

    return points_data


def make_spider_chart(
    spider_plot: matplotlib.projections.polar.PolarAxes,
    series_label: str,
    data_labels: list[str],
    data_values: list[float],
    ticks: list[float],
    tick_labels: list[str],
    spider_color: str,
    spider_marker: str,
    spider_linewidth: float,
    spider_alpha: float,
) -> matplotlib.projections.polar.PolarAxes:
    '''
    This function draws a spider/radar chart on a plot (Axes) for a given
    series of data (with values, labels, and formats).
    '''
    angles: list[float] = list(
        np.linspace(0, 2 * np.pi, len(data_labels), endpoint=False)
    )

    # We first want to plot the contour of the spider.
    # We repeat the first value at the end, since we want to close the
    # contour.
    angles_for_contour: np.ndarray = np.concatenate((angles, [angles[0]]))
    # data_labels_for_contour = np.concatenate((data_labels, [data_labels[0]]))
    data_values_for_contour: list[float] = list(
        np.concatenate((data_values, [data_values[0]]))
    )

    spider_plot.plot(
        angles_for_contour,
        data_values_for_contour,
        marker=spider_marker,
        linewidth=spider_linewidth,
        color=spider_color,
        label=series_label,
    )

    # For the fill, we use the original lists of angles and values
    spider_plot.fill(
        angles, data_values, alpha=spider_alpha, color=spider_color
    )

    spider_plot.set_thetagrids(np.array(angles) * 180 / np.pi, data_labels)
    spider_plot.set_yticks(ticks)
    spider_plot.set_yticklabels(tick_labels)
    spider_plot.legend()

    return spider_plot


def update_database_table(
    database_to_update: str,
    table_to_update: str,
    columns_to_update: list[str],
    new_values: list[ty.Any],
    query_filter_quantities: list[str],
    query_filter_types: list[str],
    query_filter_values: list[str],
) -> None:
    '''
    This function updates the values
    of one row of a table in a database.
    If you want to change multiple rows (with
    a different value for each row), then you need to iterate over the rows.
    The input parameters are:
    - database_to_update: The database you want to update (an sqlite3 file)
    - table to update: The table we want to change
    - columns to update: a list of quantities to change (column headers). Note
    that if one of the elements has a space, then it needs double quoting:
    [..., f'"{My column name with spaces}"', ...]
    - new values: a list of values (one per column to update). This function
    creates the query to update one row. To update more rows, iterate.
    - query_filter_quntities: A list of strings each representing a column
    name the user wants to filter. Again, names with spaces require
    f strings and double quotes, so add:
    f'"Surveyed Area"' to your list of filter names
    - query_filter_types: This list (that has to be the same length as the
    above liste of quantities)
    says which filter to use. Currently supported options are:
        - '='       (equal to)
        - '<'       (smaller than)
        - '>'       (larger than)
        - '!='      (not equal)
        - '<>'      (not equal)
        - '<='      (smalller or equal)
        - '>='      (larger or equal)
        - 'like'      (matches/ searches for a pattern)
        - 'between'   (between two  values)
        - 'in'        (to select multiple values for one or several columns)
    - query_filter_values: The comparison values used for the filter.
    The three special cases are:
        1) Like: This needs to be a double quote string (since it will be
        nested into a single-quote string) with percentage signs,
        such as '"%2020-05-08%"' for timestamps for May 8th, 2020
        2) Between: Provide the two  values  into a
        list. If the values arte strings that contain spaces,
        you need nested quotes, such as:
        ['"2020-05-08 00:00:00"','"2020-06-26 16:00:00"']
        3) In provide the two tuple values into a list.,
        e.g: [(52.1,4.9),(52.0,5.1)]

    '''
    first_set: bool = True
    set_query: str = ''
    for set_element, element_values in zip(columns_to_update, new_values):
        if first_set:
            set_query += 'set '
            first_set = False
        else:
            set_query += ', '

        set_query += f'{set_element} = {element_values}'

    query_filter: str = make_query_filter(
        query_filter_quantities, query_filter_types, query_filter_values
    )

    update_query: str = (
        f'update {table_to_update} ' f'{set_query} ' f'{query_filter};'
    )

    with sqlite3.connect(database_to_update) as database_connection:
        update_cursor: sqlite3.Cursor = database_connection.cursor()
        update_cursor.execute(update_query)
        database_connection.commit()


def make_query_filter(
    query_filter_quantities: list[str],
    query_filter_types: list[str],
    query_filter_values: list[str],
) -> str:
    '''
    Returns a query filter stringthat can be used in an SQL query.
     The input parameters are:
    - query_filter_quntities: A list of strings each representing a column
    name the user wants to filter. Again, names with spaces require
    f strings and double quotes, so add:
    f'"Surveyed Area"' to your list of filter names
    - query_filter_types: This list (that has to be the same length as the
    above liste of quantities)
    says which filter to use. Currently supported options are:
        - '='       (equal to)
        - '<'       (smaller than)
        - '>'       (larger than)
        - '!='      (not equal)
        - '<>'      (not equal)
        - '<='      (smalller or equal)
        - '>='      (larger or equal)
        - 'like'      (matches/ searches for a pattern)
        - 'between'   (between two  values)
        - 'in'        (to select multiple values for one or several columns)
    - query_filter_values: The comparison values used for the filter.
    The three special cases are:
        1) Like: This needs to be a double quote string (since it will be
        nested into a single-quote string) with percentage signs,
        such as '"%2020-05-08%"' for timestamps for May 8th, 2020
        2) Between: Provide the two  values  into a
        list. If the values arte strings that contain spaces,
        you need nested quotes, such as:
        ['"2020-05-08 00:00:00"','"2020-06-26 16:00:00"']
        3) In provide the two tuple values into a list.,
        e.g: [(52.1,4.9),(52.0,5.1)]
    '''
    first_filter: bool = True
    query_filter: str = ''
    for filter_quantity, filter_type, filter_value in zip(
        query_filter_quantities, query_filter_types, query_filter_values
    ):
        if first_filter:
            query_filter = 'where '
            first_filter = False
        else:
            query_filter = f'{query_filter} and'

        if filter_type.lower() == 'between':
            query_filter = (
                f'{query_filter} {filter_quantity} between {filter_value[0]} '
                f'and {filter_value[1]}'
            )
        elif filter_type.lower() == 'in':
            # We need the filter to be a string without (single) quotes
            # between brackets and the syntax and procedure are
            # different for tuples

            if type(filter_quantity) is tuple:
                tuple_content_string = ','.join(filter_quantity)
                filter_quantity = f'({tuple_content_string})'

                query_filter = (
                    f'{query_filter} {filter_quantity} in (values '
                    f'{filter_value[0]},{filter_value[1]})'
                )
            else:
                # To make a string with commas, we go through a list
                filter_value_list = [
                    f'{my_value}' for my_value in filter_value
                ]
                filter_value = ','.join(filter_value_list)
                query_filter = (
                    f'{query_filter} {filter_quantity} ' f'in ({filter_value})'
                )

        else:
            query_filter = (
                f'{query_filter} {filter_quantity} '
                f'{filter_type} {filter_value}'
            )

    return query_filter


def read_table_from_database(
    table_name: str, database_file: str
) -> pd.DataFrame:
    '''
    Reads a table from an sqlite3 database and returns it as a
    dataframe
    '''
    with sqlite3.connect(database_file) as sql_connection:
        table_query: str = read_query_generator(
            '*', f'"{table_name}"', [], [], []
        )

        table_to_read: pd.DataFrame = pd.read_sql(table_query, sql_connection)

    return table_to_read


def put_dataframe_in_word_document(
    dataframe_to_put: pd.DataFrame,
    word_document_name: str,
    number_formats: list[str] = ['.2f'],
    table_style: str = 'Normal Table',
    empty_code: str = '',
    cell_font_size: int = 11,
    headers_font_size: int = 11,
    merge_headers: bool = True,
    flip_merged_rows: bool = True,
    bottom_to_top: bool = True,
) -> None:
    '''
    This function puts a Dataframe into a given Word document.
    If the document does not exist, it is created.
    You can also optionally specify number formats for each column (in a list).
    If you do not provide the same amount of formats as there are columns
    in your DataFrame, then only the first one will be used across all columns.
    If you don't provide any format list, then the default is that your numbers
    will be displayed with two decimals.
    You can also provide a table format.
    Note that the table style must exist in the document for you to be able to
    use it (so make document, put a table in it with the style you
    want, and delete the table before saving). If it does not, you can use the
    defaults listed here:
    https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
    or simply omit the style argument.
    You can also indicate which code you want to use for empty values
    (the default is an empty string)
    Identical headers are merged by default and merged rows have their text
    flipped by default (you can also chnage the default
    bottom to top flip)
    '''

    headers_font_size = docx.shared.Pt(headers_font_size)
    cell_font_size = docx.shared.Pt(cell_font_size)
    # We first check if the fle exists. If not, we create it
    if not os.path.isfile(word_document_name):
        target_document = docx.Document()
    else:
        target_document = docx.Document(word_document_name)

    # If the user has provided less number formats (e.g. only one) than there
    # are columns (more also triggers this),
    # we put the same number format for all columns. This makes most sense if
    # there is only one format. If the user has provided more than one, but
    # less than the amount of columns, then there is an issue.
    if len(number_formats) != len(dataframe_to_put.columns):
        number_formats = [number_formats[0]] * len(dataframe_to_put.columns)
        if len(number_formats) > 1:
            print('You have provided several number formats.')
            print('But less than the amounts of columns.')
            print('We have used the first of your number formats.')
            print('Please correct your entry')

    # Weput a space before the table
    target_document.add_paragraph()

    # We want to know how many rows are for column headers and how many
    # columns are for index/row headers
    column_index_depth: int = dataframe_to_put.columns.nlevels
    row_index_depth: int = dataframe_to_put.index.nlevels

    # With these, we know how big our table has to be.
    # We create a table in the document
    table_in_document: docx.document.Table = target_document.add_table(
        rows=dataframe_to_put.shape[0] + column_index_depth,
        cols=dataframe_to_put.shape[1] + row_index_depth,
        style=table_style,
    )

    # We put the column headers in
    for column_index, column_header in enumerate(dataframe_to_put.columns):
        # If there are several levels, we need to go through them
        if column_index_depth > 1:
            for column_entry_index, column_entry in enumerate(column_header):
                current_cell = table_in_document.cell(
                    column_entry_index, column_index + row_index_depth
                )
                current_cell.text = str(
                    column_entry
                )  # We need a string to put in the table
                current_cell.paragraphs[0].runs[
                    0
                ].font.size = headers_font_size

                if merge_headers:
                    if column_index > 0:
                        previous_cell = table_in_document.cell(
                            column_entry_index,
                            column_index + row_index_depth - 1,
                        )
                        if str(column_entry) == previous_cell.text:
                            previous_cell.merge(current_cell)
                            previous_cell.text = str(column_entry)

        # Otherwise, we just put the entries
        else:
            current_cell = table_in_document.cell(
                0, column_index + row_index_depth
            )
            current_cell.text = str(
                column_header
            )  # We need a string to put in the table
            current_cell.paragraphs[0].runs[0].font.size = headers_font_size

    # We put the index/row headers in
    for row_index, row_header in enumerate(dataframe_to_put.index):
        # If there are several levels, we need to go through them
        if row_index_depth > 1:
            for row_entry_index, row_entry in enumerate(row_header):
                current_cell = table_in_document.cell(
                    row_index + column_index_depth, row_entry_index
                )
                current_cell.text = str(
                    row_entry
                )  # We need a string to put in the table
                current_cell.paragraphs[0].runs[
                    0
                ].font.size = headers_font_size
                if merge_headers:
                    if column_index > 0:
                        previous_cell = table_in_document.cell(
                            row_index + column_index_depth - 1, row_entry_index
                        )
                        if str(row_entry) == previous_cell.text:
                            previous_cell.merge(current_cell)
                            previous_cell.text = str(row_entry)
                            make_cell_text_vertical(
                                previous_cell, bottom_to_top
                            )

        # Otherwise, we just put the entries
        else:
            current_cell = table_in_document.cell(
                row_index + column_index_depth, 0
            )
            current_cell.text = str(
                row_header
            )  # We need a string to put in the table
            current_cell.paragraphs[0].runs[0].font.size = headers_font_size

    # We now put the values in
    for row_index, (row_header, row_values) in enumerate(
        dataframe_to_put.iterrows()
    ):
        for value_index, (value, number_format) in enumerate(
            zip(row_values.values, number_formats)
        ):
            current_cell = table_in_document.cell(
                row_index + column_index_depth,
                value_index + row_index_depth,
            )
            if value != value:
                # If the value is empty, we use a code for thatcell
                value = empty_code
            if type(value) is not str:
                value = f'{value:{number_format}}'
            current_cell.text = str(value)

            current_cell.paragraphs[0].runs[0].font.size = cell_font_size

            # else:
            #     # Otherwise, we use the code for empty values
            #     current_cell.text = str(empty_code)
            #     current_cell.paragraphs[0].runs[0].font.size=cell_font_size

    # We add space after the table
    target_document.add_paragraph()

    # We save the document
    target_document.save(word_document_name)


def make_cell_text_vertical(
    cell: docx.table._Cell, bottom_to_top: bool = True
):
    '''
    Changes the orientation of a Word table cell to vertical,
    with the option to get text from bottom to top (default)
    or top to bottom (set the optional bottom_to_top argument to True)
    See
    https://stackoverflow.com/questions/47738013/how-to-rotate-text-in-table-cells
    for a breakdown
    '''

    if bottom_to_top:
        orientation_code = 'btLr'
    else:
        orientation_code = 'tbRl'

    # We get the cell properties
    cell_properties = cell._tc.get_or_add_tcPr()
    # We get the textDirection and set it to our chosen direcion
    textDirection = docx.oxml.OxmlElement('w:textDirection')
    textDirection.set(docx.oxml.ns.qn('w:val'), orientation_code)
    # We change the cell'sproperties
    cell_properties.append(textDirection)


def delete_word_element(element_reference) -> None:
    '''
    Deletes a given element in a Word document.
    '''
    element = element_reference._element
    element.getparent().remove(element)


def clear_word_document(document_file_name: str) -> None:
    '''
    Clears a Word document of its elements
    (text/paragraphs, tables, pictures.)
    '''
    target_document = docx.Document(document_file_name)

    for paragraph in target_document.paragraphs:
        delete_word_element(paragraph)
    for table in target_document.tables:
        delete_word_element(table)
    for shape in target_document.inline_shapes:
        # This includes pictures
        delete_word_element(shape)

    target_document.save(document_file_name)


def get_rgb_255_code_string(
    color_name: str, color_definitions: box.Box
) -> str:
    '''
    Creates a string with rgb values (0-255),
    rgb(111, 233, 66)
    This is used for plotly.
    '''
    # We get the RGB (0-1) from the color name
    rgb_1: list[float] = get_rgb_from_name(color_name, color_definitions)
    # We convert this to a 0-255 rgb list
    rgb_255: list[int] = [int(255 * rgb_value) for rgb_value in rgb_1]

    # We make it a string (
    rgb_255_list: list[str] = list(map(str, rgb_255))
    rgb_255_code_string: str = f'rgb({", ".join(rgb_255_list)})'

    return rgb_255_code_string


def get_rgba_255_code_string(
    color_name: str, opacity: float, color_definitions: box.Box
) -> str:
    '''
    Creates a string with rgba values (0-255),
    rgb(111, 233, 66, 1)
    This is used for plotly.
    '''
    # We get the RGB (0-1) from the color name
    rgb_1: list[float] = get_rgb_from_name(color_name, color_definitions)
    # We convert this to a 0-255 rgb list
    rgba_255: list[int | float] = [int(255 * rgb_value) for rgb_value in rgb_1]
    # We add the opacity:
    rgba_255.append(opacity)

    # We make it a string (
    rgba_string_list: list[str] = list(map(str, rgba_255))
    rgba_code_string: str = f'rgba({", ".join(rgba_string_list)})'

    return rgba_code_string


def make_quantity_map(
    quantity_display_name: str,
    plot_data: pd.DataFrame,
    map_areas: pd.DataFrame,
    quantity_plot: matplotlib.axes.Axes,
    quantity_color: str,
    map_grid_plot_parameters: box.Box,
    color_definitions: box.Box,
) -> None:
    '''
    Makes one of the quantity maps in a map grid.
    '''
    # We get some display parameters
    no_data_color: list[float] = get_rgb_from_name(
        map_grid_plot_parameters.no_data_color, color_definitions
    )
    heat_bar_map: str = quantity_color
    values_column = map_grid_plot_parameters.values_column
    plot_title_font_size: int = map_grid_plot_parameters.plot_title_font_size

    map_x_range: list[float] = map_grid_plot_parameters.map_x_range
    map_y_range: list[float] = map_grid_plot_parameters.map_y_range

    # We create a range for the values to display (for the
    # scale of the legend bar).
    values_to_plot: np.ndarray = plot_data[values_column].values

    lowest_value_to_plot: float = values_to_plot.min()
    highest_value_to_plot: float = values_to_plot.max()

    display_reference_scale: list[float] = reference_scale(
        [lowest_value_to_plot, highest_value_to_plot], 1
    )
    lowest_value_to_display: float = display_reference_scale[0]
    highest_value_to_display: float = display_reference_scale[1]
    color_bar_scale: matplotlib.colors.Normalize = matplotlib.colors.Normalize(
        vmin=lowest_value_to_display, vmax=highest_value_to_display
    )

    # We plot the areas of the geographical entities (that is, the map
    # without the data) in the no-data color
    map_areas.plot(
        ax=quantity_plot,
        facecolor=no_data_color,
        edgecolor='face',
    )

    # We plot tha data on top of the map
    plot_data.plot(
        ax=quantity_plot,
        column=values_column,
        legend=True,
        norm=color_bar_scale,
        cmap=heat_bar_map,
        antialiased=True,
        edgecolor='face',
    )

    # We set the display area, remove the axes and set the title
    quantity_plot.set_ylim(map_y_range[0], map_y_range[1])
    quantity_plot.set_xlim(map_x_range[0], map_x_range[1])
    quantity_plot.axis('off')
    quantity_plot.set_title(
        quantity_display_name, fontsize=plot_title_font_size
    )


def map_grid(
    quantities_data: list[pd.DataFrame],
    quantity_display_names: list[str],
    quantity_colors: list[str],
    output_folder: str,
    map_grid_plot_parameters: box.Box,
    color_bar_definitions: box.Box,
    color_definitions: box.Box,
    dpi_to_use: int,
    file_formats: box.Box,
) -> None:
    '''
    This function creates a grid of maps. You need to give it the data you want
    to plot, the names of the quantities, and their colors, as well as
    some plot parameters (in your general parameters file, under
    a [map_grid_plot]  header). You also need to have a map areas data file
    such as this one:
    https://www.naturalearthdata.com/http//www.naturalearthdata.com/download/110m/cultural/ne_110m_admin_0_countries.zip
    You also need to provide a csv file that translates the names of the
    countries you are using into ISOA3 codes, which can be found here
    https://en.wikipedia.org/wiki/ISO_3166-1_alpha-3

    '''

    # We read some parameters

    isoA3_file: str = map_grid_plot_parameters.isoA3_file
    isoA3_codes: pd.DataFrame = pd.read_csv(f'{isoA3_file}')
    isoA3_dict: dict[str, str] = dict(
        zip(isoA3_codes.Country, isoA3_codes.IsoA3)
    )
    iso_A3_header: str = map_grid_plot_parameters.iso_A3_header
    iso_A3_header_in_map_data: str = (
        map_grid_plot_parameters.iso_A3_header_in_map_data
    )

    figure_title: str = map_grid_plot_parameters.figure_title

    number_of_rows: int = map_grid_plot_parameters.rows
    number_of_columns: int = map_grid_plot_parameters.columns
    map_data_folder: str = map_grid_plot_parameters.map_data_folder
    map_data_file: str = map_grid_plot_parameters.map_data_file
    zero_color: str = map_grid_plot_parameters.zero_color

    # We register the color bars (one per quantity color) in addition to ones
    # already existing
    for quantity_color in quantity_colors:
        color_bar_definitions[quantity_color] = [zero_color, quantity_color]
    register_color_bars(color_bar_definitions, color_definitions)

    # We read the map data from a file
    map_areas: gpd.GeoDataFrame = gpd.read_file(
        f'{map_data_folder}/{map_data_file}'
    )

    # We create a figure with one plot (grid element) for each quantity
    # we want to display
    grid_figure, quantity_plots = plt.subplots(
        number_of_rows, number_of_columns
    )

    # We iterate through the quantities (data, display name, color)
    for quantity_index, (
        quantity_data,
        quantity_display_name,
        quantity_color,
    ) in enumerate(
        zip(
            quantities_data,
            quantity_display_names,
            quantity_colors,
        )
    ):
        # We determine the row and column where the quantity plot will go
        quantity_row: int = quantity_index // number_of_columns
        quantity_column: int = quantity_index % number_of_columns
        quantity_plot: matplotlib.axes.Axes = quantity_plots[quantity_row][
            quantity_column
        ]

        # We remap the country name in the data to its ISO A3 code
        quantity_data[iso_A3_header] = quantity_data['Country'].map(isoA3_dict)
        # We create the plot data
        plot_data: pd.DataFrame = pd.merge(
            map_areas,
            quantity_data,
            left_on=iso_A3_header_in_map_data,
            right_on=iso_A3_header,
        )
        # We make the plot
        make_quantity_map(
            quantity_display_name,
            plot_data,
            map_areas,
            quantity_plot,
            quantity_color,
            map_grid_plot_parameters,
            color_definitions,
        )
    # We put a suptitle and save the figure
    grid_figure.suptitle(f'{figure_title}')
    grid_figure.tight_layout()

    save_figure(
        grid_figure, f'{figure_title}', output_folder, dpi_to_use, file_formats
    )


def put_plots_on_map(
    map_figure: matplotlib.figure.Figure,
    map_data: pd.DataFrame,
    map_parameters: dict,
    plot_y_total_values: pd.DataFrame,
    projection_type: ty.Optional[str] = None,
) -> dict[str, matplotlib.axes.Axes]:
    '''
    Puts plots/axes on a map figure. You can then draw in these.
    The plot_y_total_values are the sizes of the plots per country (for example
    the size of the stacked bar for a stacked bar plot).
    '''

    location_code_header: str = map_parameters['location_code_header']
    map_data = map_data.set_index(location_code_header)

    # We get the latitudes and longitudes of the locations (countries, e.g.)
    location_longitudes_header: str = map_parameters[
        'location_longitudes_header'
    ]
    location_latitudes_header: str = map_parameters[
        'location_latitudes_header'
    ]
    map_data['latitude'] = map_data[location_latitudes_header].values
    map_data['longitude'] = map_data[location_longitudes_header].values

    scaling_parameters: dict = map_parameters['scaling_parameters']
    # These parameters determine the size and location parameters to place
    # the plots and the necessary scaling factors. These should be adapted
    # if the plots don't come at the right place (they can change
    # if the scope of your map, or the size of your figure, for example)
    # The following values were used when making the first example using this
    # function
    # x_size = 0.005
    # y_size_max = 4
    # y_size_scale = 0.05
    # x_start = 0.51
    # y_start = 0.5
    # longitude_scaling = 0.7
    # latitude_scaling = 0.47
    # maximum_longitude = 180
    # maximum_latitude = 90
    x_size: float = scaling_parameters['x_size']
    y_size_max: float = scaling_parameters['y_size_max']
    y_size_scale: float = scaling_parameters['y_size_scale']
    x_start: float = scaling_parameters['x_start']
    y_start: float = scaling_parameters['y_start']
    longitude_scaling: float = scaling_parameters['longitude_scaling']
    latitude_scaling: float = scaling_parameters['latitude_scaling']
    maximum_longitude: float = scaling_parameters['maximum_longitude']
    maximum_latitude: float = scaling_parameters['maximum_latitude']

    # We create a dictionary that contains a plot/axis on top for each
    # location
    plots_on_top: dict[str, matplotlib.axes.Axes] = {}

    # We iterate through the locations
    for location in map_data.index:

        # We check that there are values to plot for the location
        if location in plot_y_total_values.index:
            latitude: float = map_data.loc[location]['latitude']
            longitude: float = map_data.loc[location]['longitude']
            y_size = (
                plot_y_total_values.loc[location] * y_size_scale / y_size_max
            )
            # We create the plot/Axes with the right size and scaling factors
            plot_rectangle: tuple[float, float, float, float] = (
                x_start
                * (1 + longitude_scaling * longitude / maximum_longitude),
                y_start * (1 + latitude_scaling * latitude / maximum_latitude),
                x_size,
                y_size,
            )
            plots_on_top[location] = map_figure.add_axes(
                plot_rectangle,
                projection=projection_type,
            )
    return plots_on_top


def rgba_code_color(color_rgb: tuple[int, ...], color_opacity: float) -> str:
    '''
    Gets an RGBA string from a color RGB tuple.
    This is useful for plotly.
    The A part is the color opacity.
    '''
    rgba_string: str = (
        f'rgba('
        f'{color_rgb[0]},'
        f'{color_rgb[1]},'
        f'{color_rgb[2]},'
        f'{color_opacity})'
    )
    return rgba_string


def make_sankey(
    nodes: pd.DataFrame,
    links: pd.DataFrame,
    sankey_title: str,
    output_folder: str,
    Sankey_parameters: box.Box,
    color_definitions: box.Box,
) -> None:
    '''
    Makes a Sankey plot in plotly (comes out as an html file).
    The nodes and links are in a DataFrame
    '''

    node_parameters: box.Box = Sankey_parameters.nodes

    label_padding: int = node_parameters.label_padding
    label_alignement: str = node_parameters.label_alignement

    node_labels_names_only: list[str] = pd.Series(nodes.Label).to_list()
    display_values: bool = node_parameters.display_values
    if display_values:
        values_to_add: pd.Series[float] = nodes.Value
        unit: str = node_parameters.unit
        node_labels: list[str] = []
        for node_label, value_to_add in zip(
            node_labels_names_only, values_to_add
        ):
            node_labels.append(f'{node_label}<br> {value_to_add}<br> {unit}')
    else:
        node_labels = node_labels_names_only

    node_x_positions: pd.Series[float] = nodes['X position']
    node_y_positions: pd.Series[float] = nodes['Y position']
    node_colors: pd.Series[str] = nodes.Color
    node_color_dict: dict[str, str] = dict(
        zip(node_labels_names_only, node_colors)
    )

    node_colors_rgba_codes: list[str] = []
    for node_color in node_colors:
        color_opacity: float = 1
        color_rgb: tuple[int, ...] = tuple(color_definitions[node_color])
        node_colors_rgba_codes.append(
            rgba_code_color(color_rgb, color_opacity)
        )

    link_parameters: box.Box = Sankey_parameters.links

    link_sources: pd.Series = links.Source
    link_source_indices: list[int] = [
        node_labels_names_only[node_labels_names_only == source].index[0]
        for source in link_sources
    ]

    link_targets: pd.Series = links.Target
    link_target_indices: list[int] = [
        node_labels_names_only[node_labels_names_only == target].index[0]
        for target in link_targets
    ]
    value_scaling_factor: float = link_parameters.value_scaling_factor
    link_values_unscaled: pd.Series = links.Value
    link_values: list[float] = [
        link_value / value_scaling_factor
        for link_value in link_values_unscaled
    ]
    link_colors: pd.Series[str] = links.Color
    link_opacities: pd.Series[float] = links.Opacity
    link_labels: pd.Series[str] = links.Label
    link_colors_rgba_codes: list[str] = []
    for link_color, link_opacity, source, target, link_label in zip(
        link_colors,
        link_opacities,
        link_sources,
        link_targets,
        link_labels,
    ):

        if link_color == 'source':
            link_color = node_color_dict[source]
        elif link_color == 'target':
            link_color = node_color_dict[target]

        color_rgb = tuple(color_definitions[link_color])
        link_colors_rgba_codes.append(rgba_code_color(color_rgb, link_opacity))

    sankey_figure: go.Figure = go.Figure(
        go.Sankey(
            # arrangement='snap',
            node=dict(
                label=node_labels,
                x=node_x_positions,
                y=node_y_positions,
                color=node_colors_rgba_codes,
                pad=label_padding,
                align=label_alignement,
            ),
            link=dict(
                source=link_source_indices,
                target=link_target_indices,
                value=link_values,
                color=link_colors_rgba_codes,
                label=link_labels,
            ),
        )
    )
    title_size: int = Sankey_parameters.title_size
    sankey_figure.update_layout(
        title=dict(
            text=f'{sankey_title}',
            font=dict(size=title_size),
            automargin=True,
            yref='container',
        )
    )
    sankey_figure.write_html(f'{output_folder}/{sankey_title}.html')


def get_nested_value(dictionary: dict, key_list: list[str]) -> ty.Any:
    '''
    If you give a dictionary (for example a TOML configuration file)
    and a list of nested keys, this returns the desired value.
    '''

    for key in key_list[:-1]:
        dictionary = dictionary.setdefault(key, {})
    nested_value = dictionary[key_list[-1]]

    return nested_value


def set_nested_value(
    dictionary: dict, key_list: list[str], value_to_set: ty.Any
) -> None:
    '''
    This sets the value of a nested element of a dictionary (for example
    coming from a TOML configuration file)
    '''
    for key in key_list[:-1]:
        dictionary = dictionary.setdefault(key, {})
    dictionary[key_list[-1]] = value_to_set


def function_timer(function_to_time: ty.Callable) -> ty.Callable:
    @functools.wraps(function_to_time)
    def time_wrapper(
        *function_arguments: ty.Any, **function_keywaord_arguments: ty.Any
    ) -> ty.Any:
        timer_start: float = time.perf_counter()
        function_result: ty.Any = function_to_time(
            *function_arguments, **function_keywaord_arguments
        )
        timer_end: float = time.perf_counter()
        function_run_time: float = timer_end - timer_start
        print(
            f'{function_to_time.__name__} took '
            f'{function_run_time:.2f} seconds'
        )
        return function_result

    return time_wrapper


def make_plot_sliders_dashboard(
    dashboard_parameters: box.Box,
    values_computing_function: collections.abc.Callable,
    plotting_function: collections.abc.Callable,
) -> None:
    '''
    Create a dashboard that shows a plot and sliders that can be used
    to update the plot.
    This dashboard is seen by entering http://127.0.0.1:8050/
    in your web browser.
    The parameters are in a toml file.
    The starting values of the elements that are used to compute the functions
    (including the ones in the sliders) are under the [variables]
    header. These will be update by the sliders.
    The [display] header contains the size of the plot (its
    height: The width is the golden ratio times the height) and the title
    of the dashboard.

    The [sliders] header contains the definition of the sliders:
    Here is an example (replace the values and 'my_silder')
    [sliders.my_slider]
    display_name = 'Midpoint electrity'
    id = 'midpoint_electricity'
    minimum =  2025
    maximum = 2050
    step = 1
    start_value = 2035
    ticks = [2025, 2030, 2035, 2040, 2045, 2050]
    key = 'mid_year_electric'

    '''

    # We start with getting the plotting values (y-values)
    plotting_values: list[list[float]] = values_computing_function(
        dashboard_parameters
    )
    # We create a plotly plot/figure (Dash needs this type of plot/figure).
    display_plot: plotly.graph_objs._figure.Figure = plotting_function(
        plotting_values, dashboard_parameters
    )

    # We create a dashboard
    dashboard: dash.Dash = dash.Dash(__name__)

    # We create a title
    dashboard_title: dash.html.H1 = dash.html.H1(
        dashboard_parameters.display.title, style={'textAlign': 'center'}
    )

    plot_height: float = dashboard_parameters.display.plot_height
    GOLDEN_RATIO: float = (1 + math.sqrt(5)) / 2
    plot_width: float = GOLDEN_RATIO * plot_height

    # We create a Div to display the chart
    demand_plot_display: dash.html.Div = dash.html.Div(
        children=[
            dash.html.H2(
                'Display plot',
                style={'textAlign': 'center'},
            ),
            dash.dcc.Graph(
                id='Display plot',
                figure=display_plot,
                style={
                    'width': f'{plot_width}vw',
                    'height': f'{plot_height}vh',
                },
            ),
        ]
    )

    # We create sliders
    sliders: list[dash.dcc.Slider | dash.html.H1] = []

    slider_definitions: box.Box = dashboard_parameters.sliders
    for slider in slider_definitions:
        slider_marks: dict = {
            tick: {'label': tick} for tick in slider_definitions[slider].ticks
        }
        sliders.append(
            dash.html.H1(
                slider_definitions[slider].display_name,
                style={'textAlign': 'center'},
            )
        )
        sliders.append(
            dash.dcc.Slider(
                min=slider_definitions[slider].minimum,
                max=slider_definitions[slider].maximum,
                step=slider_definitions[slider].step,
                id=slider_definitions[slider].id,
                value=slider_definitions[slider].start_value,
                marks=slider_marks,
            )
        )

    # We put all this in the layout
    dashboard.layout = dash.html.Div(
        children=[
            dashboard_title,
            demand_plot_display,
            *sliders,
            # # Need to unpack to go in children list
        ]
    )

    # We create a callback to update the plots if the inputs change
    # Below the callback , you have a function that does the updates.
    # The callback first contains the outputs (in the order
    # they appear in the return statement of the associated function).
    # Each output has two arguments: its id and its type ('figure', in this
    # case).
    # Below the outputs, you have the inputs of the updating function,
    # in the order they are listed in the function arguments.
    # Again, these require an is and a type ('value', in this case).
    # We create the elements of the callback into a list that we convert to
    # a tuple and unpack into the callback.
    # We do this (instead of putting the arguments directly into the callabck)
    # so that we can turn sliders on and off (in the parameters file, by
    # commenting them out (don't forget to do this both for the slider name
    # AND its display name)). If the quantity is not in the slider list,
    # we then can simply use the value eneterd in the parameters file.

    callback_arguments: list[dash.Input | dash.Output] = []
    # We first add the output (the figure/plot)
    callback_arguments.append(dash.Output('Display plot', 'figure'))

    # We then list the parameters that are modified
    modified_parameters: list[str] = []

    # We modify the parameers based on the slider values
    for slider in slider_definitions:
        # We add the slider value
        callback_arguments.append(
            dash.Input(slider_definitions[slider].id, 'value')
        )
        # And we add a key to find the modified parameter
        modified_parameters.append(slider_definitions[slider].key)

    # We can now perform the update
    @dashboard.callback(*callback_arguments)
    def update_plot(*callback_arguments) -> plotly.graph_objs._figure.Figure:

        for argument_index, (updated_value, key) in enumerate(
            zip(callback_arguments, modified_parameters)
        ):

            dashboard_parameters.variables[key] = updated_value

        # We remake the plot
        # We recomput the plotting values (y-values)
        plotting_values: list[list[float]] = values_computing_function(
            dashboard_parameters
        )
        # We create a plotly plot/figure (Dash needs this type of plot/figure).
        display_plot: plotly.graph_objs._figure.Figure = plotting_function(
            plotting_values, dashboard_parameters
        )
        return display_plot

    # We run the server
    dashboard.run(debug=False)


if __name__ == '__main__':

    parameters_file_name = 'test.toml'
    parameters: box.Box = parameters_from_TOML(parameters_file_name)
    color_definitions: box.Box = parameters.colors
    extra_colors: pd.DataFrame = get_extra_colors(
        color_definitions=color_definitions
    )
    color_bar_definitions: box.Box = parameters.color_bars
    register_color_bars(color_bar_definitions, color_definitions)
    # dataframe_from_Excel_table_name('Test_Table', 'Standard_Excel.xlsm')
    # check if WB and table exist. If not, return empty DF (with error message
    #  (in it?))
    # or throw an error
